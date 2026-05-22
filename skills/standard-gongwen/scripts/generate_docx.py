#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""Generate Chinese official-document DOCX files using the standard-gongwen rules."""

import argparse
import json
from pathlib import Path
from typing import Dict, List

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


class OfficialDocumentGenerator:
    FONT_SONGTI = "宋体"
    FONT_FANGSONG = "仿宋_GB2312"
    FONT_HEITI = "黑体"
    FONT_KAITI = "楷体_GB2312"
    FONT_XIAOBIAOSONG = "方正小标宋_GBK"

    def __init__(self):
        self.doc = Document()
        self._setup_page()
        self._setup_styles()

    def _setup_page(self):
        self.doc.settings.odd_and_even_pages_header_footer = True
        for section in self.doc.sections:
            section.start_type = WD_SECTION.NEW_PAGE
            section.page_width = Cm(21)
            section.page_height = Cm(29.7)
            section.top_margin = Cm(3.7)
            section.bottom_margin = Cm(3.5)
            section.left_margin = Cm(2.8)
            section.right_margin = Cm(2.6)
            section.header_distance = Cm(1.5)
            section.footer_distance = Cm(2.45)
            section.different_first_page_header_footer = True
            self._setup_page_number(section)

    def _setup_styles(self):
        style = self.doc.styles["Normal"]
        style.font.name = self.FONT_FANGSONG
        style.font.size = Pt(16)
        self._set_rfonts(style._element.rPr, self.FONT_FANGSONG)
        pf = style.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        pf.line_spacing = Pt(29)

    def _set_rfonts(self, r_pr, font_name):
        r_fonts = r_pr.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.append(r_fonts)
        for bucket in ("eastAsia", "ascii", "hAnsi", "cs"):
            r_fonts.set(qn(f"w:{bucket}"), font_name)

    def _set_run_font(self, run, font_name: str, font_size: int, bold: bool = False, color: str = "000000"):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.font.bold = bold
        self._set_rfonts(run._element.get_or_add_rPr(), font_name)
        if color != "000000":
            run.font.color.rgb = RGBColor.from_string(color)

    def _set_first_line_chars(self, paragraph, chars=200):
        p_pr = paragraph._p.get_or_add_pPr()
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            ind = OxmlElement("w:ind")
            p_pr.append(ind)
        ind.attrib.pop(qn("w:firstLine"), None)
        ind.attrib.pop(qn("w:hanging"), None)
        ind.set(qn("w:firstLineChars"), str(chars))

    def _format_paragraph(
        self,
        paragraph,
        *,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        first_line=True,
        line_spacing=29,
        line_rule=WD_LINE_SPACING.EXACTLY,
    ):
        paragraph.alignment = alignment
        pf = paragraph.paragraph_format
        pf.space_before = Pt(0)
        pf.space_after = Pt(0)
        pf.line_spacing_rule = line_rule
        if line_spacing is not None:
            pf.line_spacing = Pt(line_spacing)
        if first_line:
            self._set_first_line_chars(paragraph)

    def _add_paragraph(
        self,
        text: str = "",
        *,
        font_name: str = None,
        font_size: int = 16,
        alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
        bold: bool = False,
        first_line: bool = True,
        color: str = "000000",
        line_spacing=29,
        line_rule=WD_LINE_SPACING.EXACTLY,
    ):
        paragraph = self.doc.add_paragraph()
        self._format_paragraph(
            paragraph,
            alignment=alignment,
            first_line=first_line,
            line_spacing=line_spacing,
            line_rule=line_rule,
        )
        if text:
            run = paragraph.add_run(text)
            self._set_run_font(run, font_name or self.FONT_FANGSONG, font_size, bold, color)
        return paragraph

    def _clear_footer(self, footer):
        element = footer._element
        for child in list(element):
            element.remove(child)

    def _append_page_field(self, paragraph):
        run = paragraph.add_run("— ")
        self._set_run_font(run, self.FONT_SONGTI, 14, False)

        fld = OxmlElement("w:fldSimple")
        fld.set(qn("w:instr"), "PAGE")
        r = OxmlElement("w:r")
        r_pr = OxmlElement("w:rPr")
        self._set_rfonts(r_pr, self.FONT_SONGTI)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), "28")
        r_pr.append(sz)
        t = OxmlElement("w:t")
        t.text = "1"
        r.append(r_pr)
        r.append(t)
        fld.append(r)
        paragraph._p.append(fld)

        run = paragraph.add_run(" —")
        self._set_run_font(run, self.FONT_SONGTI, 14, False)

    def _setup_page_number(self, section):
        for footer, alignment in (
            (section.footer, WD_ALIGN_PARAGRAPH.RIGHT),
            (section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
        ):
            self._clear_footer(footer)
            paragraph = footer.add_paragraph()
            paragraph.alignment = alignment
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
            self._append_page_field(paragraph)

    def add_doc_classification(self, classification: str, period: str = None):
        text = classification + (f"★{period}" if period else "")
        return self._add_paragraph(text, font_name=self.FONT_HEITI, first_line=False)

    def add_urgency(self, urgency: str):
        return self._add_paragraph(urgency, font_name=self.FONT_HEITI, first_line=False)

    def add_document_number(self, doc_number: str):
        paragraph = self._add_paragraph(doc_number, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line=False)
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(16)
        return paragraph

    def add_issuer_mark(self, issuer: str, is_red: bool = True):
        if issuer and "文件" not in issuer:
            issuer += "文件"
        paragraph = self._add_paragraph(
            issuer,
            font_name=self.FONT_XIAOBIAOSONG,
            font_size=26,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=True,
            first_line=False,
            color="FF0000" if is_red else "000000",
            line_spacing=None,
            line_rule=WD_LINE_SPACING.SINGLE,
        )
        paragraph.paragraph_format.space_after = Pt(16)
        return paragraph

    def add_signer(self, signer_name: str):
        paragraph = self._add_paragraph("", alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
        run = paragraph.add_run("签发人：")
        self._set_run_font(run, self.FONT_FANGSONG, 16)
        run = paragraph.add_run(signer_name)
        self._set_run_font(run, self.FONT_KAITI, 16)
        return paragraph

    def add_red_separator(self):
        paragraph = self._add_paragraph(
            "━" * 40,
            font_name=self.FONT_SONGTI,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            first_line=False,
            color="FF0000",
            line_spacing=None,
            line_rule=WD_LINE_SPACING.SINGLE,
        )
        paragraph.paragraph_format.space_before = Pt(8)
        paragraph.paragraph_format.space_after = Pt(16)
        return paragraph

    def add_title(self, title: str):
        paragraph = self._add_paragraph(
            title,
            font_name=self.FONT_XIAOBIAOSONG,
            font_size=22,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            bold=False,
            first_line=False,
            line_spacing=33,
        )
        paragraph.paragraph_format.space_before = Pt(16)
        paragraph.paragraph_format.space_after = Pt(16)
        return paragraph

    def add_recipient(self, recipient: str):
        return self._add_paragraph(f"{recipient}：", first_line=False)

    def add_body_paragraph(self, text: str):
        return self._add_paragraph(text)

    def add_heading_level1(self, text: str):
        return self._add_paragraph(
            text,
            font_name=self.FONT_HEITI,
            bold=False,
            line_spacing=None,
            line_rule=WD_LINE_SPACING.SINGLE,
        )

    def add_heading_level2(self, text: str):
        return self._add_paragraph(text, font_name=self.FONT_KAITI, bold=True)

    def add_heading_level3(self, text: str):
        text = reformat_level3_marker(text)
        return self._add_paragraph(text, font_name=self.FONT_FANGSONG, bold=True)

    def add_heading_level4(self, text: str):
        return self._add_paragraph(text, font_name=self.FONT_FANGSONG, bold=False)

    def add_attachment_note(self, attachments: List[str]):
        if not attachments:
            return
        self.doc.add_paragraph()
        if len(attachments) == 1:
            self._add_paragraph(f"附件：{attachments[0]}")
            return
        self._add_paragraph(f"附件：1. {attachments[0]}")
        for i, attachment in enumerate(attachments[1:], 2):
            self._add_paragraph(f"{i}. {attachment}")

    def add_issuer_signature(self, issuer: str, date: str = None):
        self.doc.add_paragraph()
        if issuer:
            paragraph = self._add_paragraph(issuer, alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
            paragraph.paragraph_format.right_indent = Pt(64)
        if date:
            paragraph = self._add_paragraph(date, alignment=WD_ALIGN_PARAGRAPH.RIGHT, first_line=False)
            paragraph.paragraph_format.right_indent = Pt(64)

    def add_note(self, note: str):
        return self._add_paragraph(f"（{note}）")

    def add_copy_send(self, copy_to: str):
        self._add_separator_line()
        paragraph = self._add_paragraph("", first_line=False, line_spacing=None, line_rule=WD_LINE_SPACING.SINGLE)
        run = paragraph.add_run("抄送：")
        self._set_run_font(run, self.FONT_FANGSONG, 14)
        run = paragraph.add_run(f"{copy_to}。")
        self._set_run_font(run, self.FONT_FANGSONG, 14)
        return paragraph

    def add_print_info(self, print_org: str, print_date: str):
        paragraph = self._add_paragraph("", first_line=False, line_spacing=None, line_rule=WD_LINE_SPACING.SINGLE)
        run = paragraph.add_run(print_org)
        self._set_run_font(run, self.FONT_FANGSONG, 14)
        run = paragraph.add_run(" " * 20)
        self._set_run_font(run, self.FONT_FANGSONG, 14)
        run = paragraph.add_run(f"{print_date}印发")
        self._set_run_font(run, self.FONT_FANGSONG, 14)
        self._add_separator_line()

    def _add_separator_line(self):
        paragraph = self._add_paragraph(
            "─" * 50,
            font_name=self.FONT_SONGTI,
            font_size=8,
            first_line=False,
            line_spacing=None,
            line_rule=WD_LINE_SPACING.SINGLE,
        )
        paragraph.paragraph_format.space_before = Pt(4)
        paragraph.paragraph_format.space_after = Pt(4)

    def add_closing(self, closing: str = "特此通知。"):
        return self._add_paragraph(closing)

    def save(self, filepath: str):
        self.doc.save(filepath)
        return filepath


def reformat_level3_marker(text: str) -> str:
    import re

    return re.sub(r"^(\s*\d+)\.\s*", r"\1. ", text, count=1)


def add_body_by_level(gen: OfficialDocumentGenerator, paragraphs: List[str], include_level3=True):
    for paragraph in paragraphs:
        if paragraph.startswith(("一、", "二、", "三、", "四、", "五、", "六、", "七、", "八、", "九、", "十、")):
            gen.add_heading_level1(paragraph)
        elif paragraph.startswith(("（一）", "（二）", "（三）", "（四）", "（五）", "（六）", "（七）", "（八）")):
            gen.add_heading_level2(paragraph)
        elif include_level3 and paragraph.strip().startswith(("1.", "2.", "3.", "4.", "5.", "6.", "7.", "8.", "9.")):
            gen.add_heading_level3(paragraph)
        elif paragraph.strip().startswith(("（1）", "（2）", "（3）", "（4）", "（5）")):
            gen.add_heading_level4(paragraph)
        else:
            gen.add_body_paragraph(paragraph)


def create_notice(content: Dict) -> OfficialDocumentGenerator:
    gen = OfficialDocumentGenerator()
    if content.get("classification"):
        gen.add_doc_classification(content["classification"], content.get("classification_period"))
    if content.get("urgency"):
        gen.add_urgency(content["urgency"])
    gen.add_issuer_mark(content.get("issuer", ""))
    gen.add_document_number(content.get("doc_number", ""))
    if content.get("signer"):
        gen.add_signer(content["signer"])
    gen.add_red_separator()
    gen.add_title(content.get("title", ""))
    if content.get("recipient"):
        gen.add_recipient(content["recipient"])
    add_body_by_level(gen, content.get("body", []))
    if content.get("closing"):
        gen.add_closing(content["closing"])
    if content.get("attachments"):
        gen.add_attachment_note(content["attachments"])
    gen.add_issuer_signature(content.get("issuer_signature", content.get("issuer", "")), content.get("date", ""))
    if content.get("note"):
        gen.add_note(content["note"])
    if content.get("copy_to"):
        gen.add_copy_send(content["copy_to"])
    if content.get("print_org"):
        gen.add_print_info(content.get("print_org", ""), content.get("print_date", ""))
    return gen


def create_report(content: Dict) -> OfficialDocumentGenerator:
    gen = OfficialDocumentGenerator()
    gen.add_issuer_mark(content.get("issuer", ""))
    gen.add_document_number(content.get("doc_number", ""))
    gen.add_red_separator()
    gen.add_title(content.get("title", ""))
    if content.get("recipient"):
        gen.add_recipient(content["recipient"])
    add_body_by_level(gen, content.get("body", []), include_level3=False)
    gen.add_closing(content.get("closing", "特此报告。"))
    gen.add_issuer_signature(content.get("issuer_signature", content.get("issuer", "")), content.get("date", ""))
    if content.get("copy_to"):
        gen.add_copy_send(content["copy_to"])
    return gen


def create_request(content: Dict) -> OfficialDocumentGenerator:
    gen = OfficialDocumentGenerator()
    gen.add_issuer_mark(content.get("issuer", ""))
    gen.add_document_number(content.get("doc_number", ""))
    if content.get("signer"):
        gen.add_signer(content["signer"])
    gen.add_red_separator()
    gen.add_title(content.get("title", ""))
    if content.get("recipient"):
        gen.add_recipient(content["recipient"])
    add_body_by_level(gen, content.get("body", []), include_level3=False)
    gen.add_closing(content.get("closing", "妥否，请批示。"))
    if content.get("attachments"):
        gen.add_attachment_note(content["attachments"])
    gen.add_issuer_signature(content.get("issuer_signature", content.get("issuer", "")), content.get("date", ""))
    if content.get("note"):
        gen.add_note(content["note"])
    return gen


def create_letter(content: Dict) -> OfficialDocumentGenerator:
    gen = OfficialDocumentGenerator()
    gen.add_issuer_mark(content.get("issuer", ""))
    gen.add_document_number(content.get("doc_number", ""))
    gen.add_red_separator()
    gen.add_title(content.get("title", ""))
    if content.get("recipient"):
        gen.add_recipient(content["recipient"])
    add_body_by_level(gen, content.get("body", []))
    gen.add_closing(content.get("closing", "请予研究函复。"))
    gen.add_issuer_signature(content.get("issuer_signature", content.get("issuer", "")), content.get("date", ""))
    if content.get("note"):
        gen.add_note(content["note"])
    return gen


def create_minutes(content: Dict) -> OfficialDocumentGenerator:
    gen = OfficialDocumentGenerator()
    gen.add_issuer_mark(content.get("meeting_name", "会议纪要"), is_red=True)
    gen.add_title(content.get("title", f"{content.get('meeting_name', '')}纪要"))
    if content.get("time_location"):
        gen.add_body_paragraph(content["time_location"])
    if content.get("overview"):
        gen.add_body_paragraph(content["overview"])
    if content.get("attendees"):
        gen.add_body_paragraph(f"出席：{content['attendees']}")
    if content.get("absent"):
        gen.add_body_paragraph(f"请假：{content['absent']}")
    add_body_by_level(gen, content.get("body", []))
    return gen


DOCUMENT_TYPES = {
    "通知": create_notice,
    "报告": create_report,
    "请示": create_request,
    "函": create_letter,
    "纪要": create_minutes,
}


EXAMPLE_CONTENT = {
    "issuer": "XXX公司",
    "doc_number": "XX〔2024〕1号",
    "title": "关于开展互联网服务统一管理工作的通知",
    "recipient": "各部门、各子公司",
    "body": [
        "为进一步规范集团互联网服务管理，加强网络安全防护，提升信息化管理水平，现将有关事项通知如下：",
        "一、工作目标",
        "建立健全互联网服务管理体系，消除安全隐患，保障信息系统安全稳定运行。",
        "二、主要内容",
        "（一）微信公众号",
        "包括公众号名称、账号主体、运营负责人及联系方式。",
        "（二）网站信息",
        "包括网站名称、域名、IP地址、备案号、服务器位置。",
        "1. 信息报送要求",
        "各单位应确保信息准确、完整、及时。",
    ],
    "closing": "特此通知。",
    "attachments": ["互联网服务信息登记表"],
    "issuer_signature": "XXX公司",
    "date": "2024年1月15日",
    "copy_to": "集团各部门",
    "print_org": "集团办公室",
    "print_date": "2024年1月16日",
}


def generate_document(doc_type: str, content: Dict, output_path: str) -> str:
    if doc_type not in DOCUMENT_TYPES:
        raise ValueError(f"不支持的公文类型: {doc_type}。支持的类型: {list(DOCUMENT_TYPES.keys())}")
    generator = DOCUMENT_TYPES[doc_type](content)
    generator.save(output_path)
    return output_path


def main():
    parser = argparse.ArgumentParser(description="Generate a standard-gongwen DOCX.")
    parser.add_argument("doc_type", nargs="?", choices=sorted(DOCUMENT_TYPES), default="通知")
    parser.add_argument("content_json", nargs="?", help="UTF-8 JSON file containing document content.")
    parser.add_argument("output_docx", nargs="?", default="example_notice.docx")
    parser.add_argument("--example", action="store_true", help="Generate an example notice without reading JSON.")
    args = parser.parse_args()

    if args.example or not args.content_json:
        content = EXAMPLE_CONTENT
    else:
        content = json.loads(Path(args.content_json).read_text(encoding="utf-8"))

    output = generate_document(args.doc_type, content, args.output_docx)
    print(f"公文已生成: {output}")


if __name__ == "__main__":
    main()
