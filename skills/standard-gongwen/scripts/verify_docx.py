#!/usr/bin/env python3
import argparse
from pathlib import Path

from docx import Document


def east_asia_font(run):
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return None
    return r_pr.rFonts.get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia")


def script_fonts(run):
    r_pr = run._element.rPr
    if r_pr is None or r_pr.rFonts is None:
        return {}
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    return {bucket: r_pr.rFonts.get(f"{ns}{bucket}") for bucket in ("eastAsia", "ascii", "hAnsi", "cs")}


def cm(value):
    return round(value.cm, 2)


def verify(path):
    doc = Document(str(path))
    print("opened: true")
    print(f"paragraphs: {len(doc.paragraphs)}")
    print(f"tables: {len(doc.tables)}")
    print(f"sections: {len(doc.sections)}")
    for index, section in enumerate(doc.sections):
        print(f"section_{index}_page_cm:", f"{cm(section.page_width)} x {cm(section.page_height)}")
        print(
            f"section_{index}_margins_cm:",
            cm(section.top_margin),
            cm(section.bottom_margin),
            cm(section.left_margin),
            cm(section.right_margin),
        )
        print(f"section_{index}_header_footer_cm:", cm(section.header_distance), cm(section.footer_distance))
        print(f"section_{index}_different_first_page:", bool(section.different_first_page_header_footer))
    print("odd_even_pages:", bool(doc.settings.odd_and_even_pages_header_footer))

    printed = 0
    for i, paragraph in enumerate(doc.paragraphs):
        if not paragraph.text.strip() or not paragraph.runs:
            continue
        run = paragraph.runs[0]
        size = run.font.size.pt if run.font.size is not None else None
        print(
            f"sample_paragraph_{i}:",
            paragraph.text[:60],
            "font_buckets=",
            script_fonts(run),
            "size=",
            size,
            "bold=",
            run.bold,
        )
        printed += 1
        if printed >= 10:
            break

    for table_index, table in enumerate(doc.tables[:3]):
        printed_table = False
        for row_index, row in enumerate(table.rows):
            for cell_index, cell in enumerate(row.cells):
                paragraph = next((item for item in cell.paragraphs if item.text.strip() and item.runs), None)
                if paragraph is None:
                    continue
                run = paragraph.runs[0]
                size = run.font.size.pt if run.font.size is not None else None
                print(
                    f"sample_table_{table_index}_{row_index}_{cell_index}:",
                    paragraph.text[:60],
                    "font_buckets=",
                    script_fonts(run),
                    "size=",
                    size,
                    "bold=",
                    run.bold,
                )
                printed_table = True
                break
            if printed_table:
                break


def main():
    parser = argparse.ArgumentParser(description="Inspect key 标准公文 DOCX formatting fields.")
    parser.add_argument("docx")
    args = parser.parse_args()
    verify(Path(args.docx))


if __name__ == "__main__":
    main()
