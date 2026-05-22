#!/usr/bin/env python3
import argparse
import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


BODY_FONT = "仿宋_GB2312"
PROJECT_REPORT_COVER_LABELS = (
    "项目名称",
    "项目类型",
    "项目级别",
    "项目主管部门",
)


def set_font(run, font, size_pt=None, bold=None):
    # Set every script bucket so digits, Latin text, and symbols do not fall
    # back to Times New Roman or another Western font.
    run.font.name = font
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    for bucket in ("eastAsia", "ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{bucket}"), font)


def set_first_line_chars(paragraph, chars=200):
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:hanging"), None)
    ind.set(qn("w:firstLineChars"), str(chars))


def clear_first_line_indent(paragraph):
    p_pr = paragraph._p.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        return
    ind.attrib.pop(qn("w:firstLine"), None)
    ind.attrib.pop(qn("w:firstLineChars"), None)
    ind.attrib.pop(qn("w:hanging"), None)


def apply_paragraph_format(
    paragraph,
    *,
    font=BODY_FONT,
    size=16,
    bold=False,
    alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
    line_spacing=29,
    line_rule=WD_LINE_SPACING.EXACTLY,
    first_line=True,
):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.line_spacing_rule = line_rule
    if line_spacing is not None:
        pf.line_spacing = Pt(line_spacing)
    paragraph.alignment = alignment
    if first_line:
        set_first_line_chars(paragraph, 200)
    else:
        clear_first_line_indent(paragraph)

    for run in paragraph.runs:
        set_font(run, font, size, bold)


def paragraph_kind(text):
    stripped = text.strip()
    if not stripped:
        return "empty"
    if re.match(r"^[一二三四五六七八九十]+、", stripped):
        return "level1"
    if re.match(r"^（[一二三四五六七八九十]+）", stripped):
        return "level2"
    if re.match(r"^\d+\.\s*", stripped):
        return "level3"
    if re.match(r"^（\d+）", stripped):
        return "level4"
    return "body"


def existing_alignment(paragraph, fallback=WD_ALIGN_PARAGRAPH.JUSTIFY):
    return paragraph.alignment if paragraph.alignment is not None else fallback


def normalized_text(text):
    return re.sub(r"\s+", "", text or "")


def normalize_numbered_heading_text(paragraph):
    text = paragraph.text
    new_text = re.sub(r"^(\s*\d+)\.\s*", r"\1. ", text, count=1)
    if new_text == text:
        return
    paragraph.clear()
    paragraph.add_run(new_text)


def clear_footer(footer):
    element = footer._element
    for child in list(element):
        element.remove(child)


def append_page_field(paragraph, font="宋体", size=14):
    run = paragraph.add_run("— ")
    set_font(run, font, size, False)

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    r = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    for bucket in ("eastAsia", "ascii", "hAnsi", "cs"):
        r_fonts.set(qn(f"w:{bucket}"), font)
    r_pr.append(r_fonts)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(size * 2))
    r_pr.append(sz)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(r_pr)
    r.append(t)
    fld.append(r)
    paragraph._p.append(fld)

    run = paragraph.add_run(" —")
    set_font(run, font, size, False)


def setup_page_numbers(section):
    for footer, alignment in (
        (section.footer, WD_ALIGN_PARAGRAPH.RIGHT),
        (section.even_page_footer, WD_ALIGN_PARAGRAPH.LEFT),
    ):
        clear_footer(footer)
        paragraph = footer.add_paragraph()
        paragraph.alignment = alignment
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
        append_page_field(paragraph)


def format_table_text(table, *, default_size=None, force_size=False):
    """Normalize table run fonts without changing table layout or paragraph geometry."""
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    size = default_size if force_size or run.font.size is None else run.font.size.pt
                    set_font(run, BODY_FONT, size, None)


def first_level1_index(paragraphs):
    for index, paragraph in enumerate(paragraphs):
        if paragraph_kind(paragraph.text) == "level1":
            return index
    return 0


def is_project_report_title(paragraph, text):
    if "技术论证报告" in text:
        return True
    run = next((run for run in paragraph.runs if run.text), None)
    size = run.font.size.pt if run is not None and run.font.size is not None else None
    return paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER and size is not None and size >= 20


def is_project_report_cover_label(text):
    compact = normalized_text(text)
    return any(label in compact for label in PROJECT_REPORT_COVER_LABELS)


def apply_project_report_cover_format(paragraph):
    text = paragraph.text.strip()
    if not text:
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        return
    if is_project_report_title(paragraph, text):
        apply_paragraph_format(
            paragraph,
            font="方正小标宋简体",
            size=22,
            bold=False,
            alignment=WD_ALIGN_PARAGRAPH.CENTER,
            line_spacing=33,
            first_line=False,
        )
    elif is_project_report_cover_label(text):
        apply_paragraph_format(
            paragraph,
            font=BODY_FONT,
            size=14,
            bold=True,
            alignment=existing_alignment(paragraph, WD_ALIGN_PARAGRAPH.LEFT),
            line_spacing=None,
            line_rule=WD_LINE_SPACING.ONE_POINT_FIVE,
            first_line=False,
        )
    else:
        apply_paragraph_format(
            paragraph,
            font=BODY_FONT,
            size=16,
            bold=None,
            alignment=existing_alignment(paragraph, WD_ALIGN_PARAGRAPH.CENTER),
            line_spacing=29,
            first_line=False,
        )


def apply_standard_sections(doc):
    doc.settings.odd_and_even_pages_header_footer = True
    for section in doc.sections:
        section.start_type = WD_SECTION.NEW_PAGE
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(3.7)
        section.bottom_margin = Cm(3.5)
        section.left_margin = Cm(2.8)
        section.right_margin = Cm(2.6)
        section.header_distance = Cm(1.5)
        section.footer_distance = Cm(2.45)
        section.different_first_page_header_footer = True
        setup_page_numbers(section)


def apply_project_report_sections(doc):
    doc.settings.odd_and_even_pages_header_footer = False
    for index, section in enumerate(doc.sections):
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.different_first_page_header_footer = False
        if index == 0 and len(doc.sections) > 1:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(2.5)
            section.right_margin = Cm(2.5)
            section.header_distance = Cm(1.5)
            section.footer_distance = Cm(1.75)
        else:
            section.top_margin = Cm(2.54)
            section.bottom_margin = Cm(2.54)
            section.left_margin = Cm(3.17)
            section.right_margin = Cm(3.17)
            section.header_distance = Cm(1.27)
            section.footer_distance = Cm(1.27)


def format_standard_paragraph(paragraph):
    kind = paragraph_kind(paragraph.text)
    if kind == "empty":
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
    elif kind == "level1":
        apply_paragraph_format(
            paragraph,
            font="黑体",
            size=16,
            bold=False,
            line_spacing=None,
            line_rule=WD_LINE_SPACING.SINGLE,
        )
    elif kind == "level2":
        apply_paragraph_format(paragraph, font="楷体_GB2312", size=16, bold=True)
    elif kind == "level3":
        normalize_numbered_heading_text(paragraph)
        apply_paragraph_format(paragraph, font=BODY_FONT, size=16, bold=True)
    elif kind == "level4":
        apply_paragraph_format(paragraph, font=BODY_FONT, size=16, bold=False)
    else:
        apply_paragraph_format(paragraph, font=BODY_FONT, size=16, bold=False)


def format_project_report_paragraph(paragraph):
    kind = paragraph_kind(paragraph.text)
    if kind == "empty":
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
    elif kind == "level1":
        apply_paragraph_format(
            paragraph,
            font="黑体",
            size=16,
            bold=False,
            alignment=existing_alignment(paragraph),
            line_spacing=None,
            line_rule=WD_LINE_SPACING.SINGLE,
        )
    elif kind == "level2":
        apply_paragraph_format(
            paragraph,
            font="楷体_GB2312",
            size=16,
            bold=False,
            alignment=existing_alignment(paragraph),
        )
    elif kind == "level3":
        normalize_numbered_heading_text(paragraph)
        apply_paragraph_format(
            paragraph,
            font="楷体_GB2312",
            size=16,
            bold=False,
            alignment=existing_alignment(paragraph),
        )
    else:
        apply_paragraph_format(
            paragraph,
            font=BODY_FONT,
            size=16,
            bold=False,
            alignment=existing_alignment(paragraph),
        )


def format_document(input_path, output_path=None, backup=False, profile="standard"):
    input_path = Path(input_path)
    output_path = Path(output_path) if output_path else input_path
    if backup and output_path == input_path:
        backup_path = input_path.with_name(f"{input_path.stem}.backup-before-standard-gongwen{input_path.suffix}")
        shutil.copy2(input_path, backup_path)

    doc = Document(str(input_path))
    if profile == "standard":
        apply_standard_sections(doc)
        for paragraph in doc.paragraphs:
            format_standard_paragraph(paragraph)
        for table in doc.tables:
            format_table_text(table, default_size=16, force_size=True)
    elif profile == "project-report":
        apply_project_report_sections(doc)
        first_body = first_level1_index(doc.paragraphs)
        for index, paragraph in enumerate(doc.paragraphs):
            if index < first_body:
                apply_project_report_cover_format(paragraph)
            else:
                format_project_report_paragraph(paragraph)
        for table in doc.tables:
            format_table_text(table)
    else:
        raise ValueError(f"Unsupported profile: {profile}")

    doc.save(str(output_path))


def main():
    parser = argparse.ArgumentParser(description="Apply 标准公文 formatting to a DOCX file.")
    parser.add_argument("input_docx")
    parser.add_argument("--output", help="Write to a different DOCX path. Defaults to in-place.")
    parser.add_argument("--backup", action="store_true", help="Create a sibling backup before in-place overwrite.")
    parser.add_argument(
        "--profile",
        choices=("standard", "project-report"),
        default="standard",
        help="Use standard official-document rules or the project technical report profile.",
    )
    args = parser.parse_args()
    format_document(args.input_docx, args.output, args.backup, args.profile)
    print(f"Saved formatted document: {args.output or args.input_docx}")


if __name__ == "__main__":
    main()
